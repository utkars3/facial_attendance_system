import streamlit as st
from Home import face_rec
import pandas as pd
from io import BytesIO
from docx import Document

#st.set_page_config(page_title='Reporting',layout='wide')
st.subheader('Reporting')

#Retrive logs data and show in Report.py
#extract data from redis list
name='attendance:logs'
def load_logs(name,end=-1):
    logs_list=face_rec.r.lrange(name,start=0,end=end) # extract all data from redis db
    return logs_list

#tabs to show the info
tab1,tab2,tab3=st.tabs(['Enrolled Data','Logs',"Attendance Report"])

with tab1:
    if st.button('Refresh Data'):
    #retrive data from redis db
        with st.spinner('Retriving Data from Redis DB'):
            redis_face_db=face_rec.retrive_data('academy:enrollment')
            st.dataframe(redis_face_db[['Name','Role']])
        
with tab2:
    if st.button('Refresh Logs'):
        st.write(load_logs(name=name))

with tab3:
    st.subheader('Attendance Report')       

    #load logs into attribute logs_list
    logs_list=load_logs(name=name)
    

    #step-1 convert the logs that in list of bytes into list of strings
    convert_byte_to_string=lambda x: x.decode('utf-8')
    logs_list_string=list(map(convert_byte_to_string,logs_list))

    #step-2 split string by @ and create nested list
    split_string=lambda x: x.split('@')
    logs_nested_list=list(map(split_string,logs_list_string))

    #convert nested list into dataframe
    logs_df=pd.DataFrame(logs_nested_list,columns=['Name','Role','Timestamp'])

    #step-3 Time based Analysis or Report
    logs_df['Timestamp']=pd.to_datetime(logs_df['Timestamp'])
    logs_df['Date']=logs_df['Timestamp'].dt.date

    # st.dataframe(logs_df)

    #step-3.1 Cal intime and outtime
    #In time: At which person is first detected in that date (min Timestamp of the data)
    #out time: At which person is last detected in that date (max Timestamp of the data)

    report_df=logs_df.groupby(by=['Date','Name','Role']).aggregate(
        In_time=pd.NamedAgg('Timestamp','min'), #In time
        Out_time=pd.NamedAgg('Timestamp','max') #out time
    ).reset_index()

    report_df['In_time']=pd.to_datetime(report_df['In_time'])
    report_df['Out_time']=pd.to_datetime(report_df['Out_time'])
    report_df['Duration']=report_df['Out_time'] - report_df['In_time']
    def format_duration(duration):
        if pd.isnull(duration):  # Handle missing values
            return "00:00:00"
        total_seconds = int(duration.total_seconds())
        hours, remainder = divmod(total_seconds, 3600)
        minutes, seconds = divmod(remainder, 60)
        return f"{hours:02}:{minutes:02}:{seconds:02}"

    report_df['Duration'] = report_df['Duration'].apply(format_duration)


    #step 4 : Marking student present or absent
    all_dates=report_df['Date'].unique()
    name_role=report_df[['Name','Role']].drop_duplicates().values.tolist()

    date_name_role_zip=[]

    for dt in all_dates:
        for name,role in name_role:
            date_name_role_zip.append([dt,name,role])
    
    date_name_role_zip_df=pd.DataFrame(date_name_role_zip,columns=['Date','Name','Role'])

    #left join with report_df

    date_name_role_zip_df=pd.merge(date_name_role_zip_df,report_df,how='left',on=['Date','Name','Role'])

    #Duration
    #Hours
    # date_name_role_zip_df['Duration_seconds']=pd.to_datetime(date_name_role_zip_df['Duration']).dt.seconds
    date_name_role_zip_df['Duration_seconds'] = pd.to_timedelta(date_name_role_zip_df['Duration']).dt.total_seconds()
    date_name_role_zip_df['Duration_hours']= date_name_role_zip_df['Duration_seconds']/(60*60) #hours

    def status_marker(x):
        if pd.Series(x).isnull().all():
            return 'Absent'
        
        elif x==0:
            return 'Student inside class'
        
        elif x>0 and x<1:
            return 'Present (less than 1 hr)'
        
        elif x>=1 and x<4:
            return 'Present (less than 4 hr)'
        
        elif x>=4:
            return 'Present'
        
    date_name_role_zip_df['Status']=date_name_role_zip_df['Duration_hours'].apply(status_marker)   

    #search bar
    options=['Name','Role','Date','Status','Duration_hours']
    selected_option=st.selectbox('Select filter for search',options)
    search_term=st.text_input(f'Search by {selected_option}')

    if search_term:
        if selected_option in ['Name','Role','Status']:
            data_filtered=date_name_role_zip_df[date_name_role_zip_df[selected_option].str.contains(search_term,case=False,na=False)]
        elif selected_option == 'Date':
            try:
                _search_date=pd.to_datetime(search_term).date()
                data_filtered=date_name_role_zip_df[date_name_role_zip_df['Date']==_search_date]
            except ValueError:
                st.error("Invalid Date Format. Please enter date in YYYY-MM-DD") 
                data_filtered=date_name_role_zip_df  
        elif selected_option=='Duration_hours':
            try:
                _search_duration_hour=float(search_term)
                data_filtered=date_name_role_zip_df[date_name_role_zip_df['Duration_hours']>=_search_duration_hour]
            except ValueError:
                st.error("Invalid Date Format. Please enter date in YYYY-MM-DD") 
                data_filtered=date_name_role_zip_df  

    else:
        data_filtered=date_name_role_zip_df    

   
    st.dataframe(data_filtered)




    #downloading
    def generate_excel(dataframe):
        output = BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            dataframe.to_excel(writer, index=False, sheet_name='Attendance Report')
            writer.save()
        processed_data = output.getvalue()
        return processed_data

# Function to generate Word file
    def generate_word(dataframe):
        document = Document()
        document.add_heading('Attendance Report', level=1)

        # Adding table
        table = document.add_table(rows=1, cols=len(dataframe.columns))
        table.style = 'Table Grid'

        # Add header
        header_cells = table.rows[0].cells
        for idx, column in enumerate(dataframe.columns):
            header_cells[idx].text = column

        # Add data rows
        for _, row in dataframe.iterrows():
            cells = table.add_row().cells
            for idx, value in enumerate(row):
                cells[idx].text = str(value)

        output = BytesIO()
        document.save(output)
        processed_data = output.getvalue()
        return processed_data
    

    excel_data = generate_excel(date_name_role_zip_df)
    st.download_button(
        label="Download Excel",
        data=excel_data,
        file_name="attendance_report.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

    # Word download button
    word_data = generate_word(date_name_role_zip_df)
    st.download_button(
        label="Download Word",
        data=word_data,
        file_name="attendance_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )



